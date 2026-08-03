import html
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
INDEX_HTML = REPO_ROOT / "index.html"
DEFAULT_OUTPUT = REPO_ROOT / "exports" / "幸福学院市场服务手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1D1D1F"
MUTED = "555555"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


def read_manual_data():
    text = INDEX_HTML.read_text(encoding="utf-8")
    match = re.search(
        r'<script type="application/json" id="manualData">(.*?)</script>',
        text,
        flags=re.S,
    )
    if not match:
        raise RuntimeError("Cannot find manualData in index.html")
    return json.loads(html.unescape(match.group(1)))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2


def add_title(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("幸福学院市场服务手册")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("我不相信人，我只相信流程。")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(18)
    run = note.add_run("导出自幸福驿站市场服务流程页面，适合培训、通关、复盘和打印查阅。")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_numbered_list(doc, items):
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"{index}. {item}")


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(f"- {item}")


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(5.2)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        set_cell_text(row.cells[0], label, bold=True, color=DARK_BLUE)
        set_cell_text(row.cells[1], value)
    doc.add_paragraph()


def add_stage(doc, stage):
    doc.add_heading(f'{stage["number"]}. {stage["title"]}', level=2)
    add_label_table(
        doc,
        [
            ("入口/定位", stage.get("tag", "")),
            ("目标", stage.get("goal", "")),
            ("查找来源", stage.get("source", "")),
        ],
    )
    if stage.get("process"):
        doc.add_heading("流程步骤", level=3)
        add_numbered_list(doc, stage["process"])
    if stage.get("standards"):
        doc.add_heading("执行标准", level=3)
        add_numbered_list(doc, stage["standards"])
    if stage.get("script"):
        doc.add_heading("话术原文", level=3)
        for line in stage["script"]:
            doc.add_paragraph(str(line))
    if stage.get("review"):
        doc.add_heading("复盘问题", level=3)
        add_numbered_list(doc, stage["review"])
    if stage.get("original"):
        doc.add_heading("标准原文", level=3)
        add_raw_lines(doc, stage["original"])


def add_raw_lines(doc, lines):
    for line in lines:
        text = str(line).strip()
        if not text:
            continue
        clean_heading = text.rstrip("：:")
        if clean_heading in {"心态标准", "动作标准", "学习要求", "作业要求", "复制标准", "交作业的标准"}:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(clean_heading)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
            continue
        ordered = re.match(r"^([0-9]{1,2}[.、）)]|第[0-9一二三四五六七八九十]+[天步])\s*(.+)$", text)
        if ordered:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(text)
        else:
            doc.add_paragraph(text)


def add_salon_reviews(doc, data):
    doc.add_heading("镜子库：复盘清单", level=1)
    doc.add_heading("沙龙复盘", level=2)
    add_numbered_list(doc, data.get("salonReview", []))

    doc.add_heading("7天训练营复盘", level=2)
    camp_questions = [
        "学员是不是从沙龙或者转介绍筛选来的？有没有跳过筛选？",
        "是不是一个教练带2个顾客？有没有混营？",
        "每天课程之前有没有通关？教练是否掌握正确标准？",
        "学员有没有全程开视频、提前进会议室、按要求参与？",
        "作业有没有按时提交？是否具体到对象、场景、原话、反馈和心情？",
        "不配合、不按要求实践的学员有没有按标准处理？",
        "沙龙和7天训练营搜集来的榜样，是否统一进入一个大群，等待被采访？",
    ]
    add_numbered_list(doc, camp_questions)


def add_model_collection(doc, data):
    collection = data.get("modelCollection", {})
    doc.add_heading("榜样采访", level=1)
    for index, step in enumerate(collection.get("steps", []), start=1):
        doc.add_heading(f'{index}. {step.get("title", "")}', level=2)
        doc.add_paragraph(step.get("body", ""))
        if step.get("purpose"):
            add_label_table(doc, [("核心目的", step["purpose"])])
    doc.add_heading("榜样采访问句模板", level=2)
    add_numbered_list(doc, collection.get("questions", []))


def add_frameworks(doc, data):
    doc.add_heading("市场服务", level=1)
    for item in data.get("serviceFramework", []):
        doc.add_heading(f'{item.get("number", "")}. {item.get("title", "")}', level=2)
        doc.add_paragraph(item.get("body", ""))

    doc.add_heading("总监思维", level=1)
    for index, item in enumerate(data.get("directorThinking", []), start=1):
        doc.add_heading(f'{index}. {item.get("title", "")}', level=2)
        doc.add_paragraph(item.get("body", ""))


def split_display_lines(text):
    lines = []
    for raw in str(text or "").replace("\u00a0", " ").splitlines():
        raw = raw.strip()
        if not raw:
            continue
        parts = re.split(r"(?<=[。！？；;])(?=(第[0-9一二三四五六七八九十]+[步天]|[0-9]{1,2}[.、]|✅|⚠))", raw)
        lines.extend(part.strip() for part in parts if part.strip())
    return lines


def is_wrapper_line(line, title):
    text = str(line).strip()
    bare = re.sub(r"^[0-9]{1,2}[.、）)]\s*", "", text).rstrip("：:").strip()
    wrapper_titles = {
        title,
        title.replace("模块", ""),
        "7天幸福训练营",
        "幸福早课人才培养营",
        "幸福学院市场培训手册",
        "幸福学院市场培训手册（1.0）",
        "我不相信人，我只相信流程。",
        "0）",
        "0)",
    }
    return any(text == item or bare == item or bare.startswith(item + "（") for item in wrapper_titles)


def add_raw_sources(doc, data):
    doc.add_heading("原文模块", level=1)
    for source in data.get("rawMirrorSources", []):
        title = source.get("title", "")
        doc.add_heading(title, level=2)
        lines = [line for line in split_display_lines(source.get("content", "")) if not is_wrapper_line(line, title)]
        add_raw_lines(doc, lines)


def build_docx(output_path):
    data = read_manual_data()
    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("大流程｜从陌生人到合作伙伴", level=1)
    for stage in data.get("masterFlow", []):
        add_stage(doc, stage)

    add_page_break_before(doc.add_heading("镜子库｜用流程标准照镜子", level=1))
    add_salon_reviews(doc, data)
    add_model_collection(doc, data)
    add_frameworks(doc, data)
    add_raw_sources(doc, data)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT
    build_docx(target)
    print(target)
