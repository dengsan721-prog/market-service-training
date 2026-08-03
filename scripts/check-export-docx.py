import sys
from pathlib import Path

from docx import Document


def main():
    path = Path(sys.argv[1])
    if not path.exists():
        raise SystemExit(f"Missing DOCX: {path}")
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    required = [
        "幸福学院市场服务手册",
        "我不相信人，我只相信流程。",
        "大流程｜从陌生人到合作伙伴",
        "01. 来人先承接",
        "02. 免费幸福沙龙",
        "03. 7天幸福训练营",
        "镜子库：复盘清单",
        "榜样采访",
        "市场服务",
        "总监思维",
        "原文模块",
        "沙龙模块",
        "7天训练营模块",
        "幸福早课人才培养营模块",
        "榜样选拔与教练招募",
    ]
    forbidden_exact = {
        "0）",
        "7天幸福训练营",
        "幸福早课人才培养营",
        "幸福学院市场培训手册（1.0）",
    }
    missing = [item for item in required if item not in text]
    paragraph_forbidden = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.text.strip() in forbidden_exact
    ]
    auto_numbered = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style and paragraph.style.name == "List Number"
    ]
    if missing or paragraph_forbidden or auto_numbered:
        if missing:
            print("Missing:")
            print("\n".join(missing))
        if paragraph_forbidden:
            print("Forbidden wrapper paragraphs:")
            print("\n".join(paragraph_forbidden))
        if auto_numbered:
            print("Auto-numbered paragraphs should preserve literal source numbering:")
            print("\n".join(auto_numbered[:20]))
        raise SystemExit(1)
    print(f"export docx check passed: {path}")


if __name__ == "__main__":
    main()
